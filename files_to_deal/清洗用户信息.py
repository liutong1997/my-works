import docx
from multiprocessing import Process, Value


def deal_alltable(file):
    # 读取文件

    # 读取所有表格
    tables = file.tables
    # 抽取表格数据
    for table in tables:
        # 判断是否为含有编号的单元格，含有即为用户信息表
        if table.cell(0, 0).text == '编号':
            # 创建变量m收集行数，每张表进行一次初始化
            m = 0

            # 遍历行
            for i in range(len(table.rows)):
                # 创建变量n收集列数，每行进行一次初始化
                n = 0
                # 统计行数
                m += 1
                # 每遍历一次行遍历该行中所有列
                for j in range(len(table.rows[i].cells)):
                    # 统计列数
                    n += 1
                    # 判断如果为第二行则开始清洗内容
                    if i >= 1:
                        # 清空段落数据
                        run = table.cell(i, j).paragraphs[0].clear()
            table.cell(1, 0).merge(table.cell(m - 1, n - 1))


if __name__ == '__main__':
    file = docx.Document(input('请输入要读取的正确的docx文件路径:'))
    # 调用清洗函数
    deal_alltable(file)
    file.save(input('请输入要写入的正确的docx文件路径，如果与读取路径重名将覆盖源文件，请慎重:'))
