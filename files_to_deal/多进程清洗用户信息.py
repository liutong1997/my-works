# coding = utf-8
from multiprocessing import Process
import docx
import os
import glob


def clear_all_user_information(t):
    if t.cell(0, 0).text == '编号':  # 判断是否为含有编号的单元格，含有即为用户信息表
        m = 0  # 创建变量m收集行数，每张表进行一次初始化
        # 遍历行
        for i in range(len(t.rows)):
            n = 0  # 创建变量n收集列数，每行进行一次初始化
            m += 1  # 统计行数
            # 每遍历一次行遍历该行中所有列
            for j in range(len(t.rows[i].cells)):
                n += 1  # 统计列数
                if i >= 1:  # 判断如果为第二行则开始清洗内容
                    run = t.cell(i, j).paragraphs[0].clear()  # 清空段落数据
        t.cell(1, 0).merge(t.cell(m - 1, n - 1))  # 合并单元格


if __name__ == '__main__':
    file_path = os.getcwd()
    for file_name in glob.glob(os.path.join(file_path, 'new*.docx')):
        file = docx.Document(file_name)
        tables = file.tables  # 读取所有表格
        # 遍历所有表，迭代得出每一张表
        for i in range(len(tables)):
            t = tables[i]  # 将每一张表传递给参数t
            p = Process(target=clear_all_user_information(t))  # Process 对象只是一个子任务，运行该任务时系统会自动创建一个子进程
            # 启动一个子进程来运行子任务
            p.start()
            p.join()
            # 子进程完成后，继续运行主进程,保存文件
        new_name = file_name.strip(file_path)  # 去掉文件路径只保留文件名
        file.save('hello' + new_name)  # 保存文件
