# coding = utf-8
import shutil
from multiprocessing import Process
import docx
import os
import glob


# 定义自动创建文件夹的函数
def make_folder(p):
    if os.path.exists(p):  # 判断文件夹是否存在
        shutil.rmtree(p)  # 如果存在删除原有目录的文件夹以及其中所有文件
    os.mkdir(p)


def clear_all_user_information(t):
    if t.cell(0, 0).text == '编号':  # 判断是否为含有编号的单元格，含有即为用户信息表
        m = 0  # 创建变量m收集行数，每张表进行一次初始化
        # 遍历行
        for i in range(len(t.rows)):
            n = 0  # 创建变量n收集列数，每行进行一次初始化
            m += 1  # 统计行数
            # 每遍历一次行遍历该行中所有列
            for j in range(len(t.rows[i].cells)):
                n += 1  # 统计列数
                if t.cell(i, j).text in ['政治面貌', '班级', '电子邮件', '联系电话']:  # 判断如果为第二行则开始清洗内容
                    run = t.cell(i, j).paragraphs[0].clear()  # 清空单元格段落数据
                    t.cell(i, j).merge(t.cell(i, len(t.rows[i].cells) - 1))  # 合并单元格
                    run = t.cell(i, j).paragraphs[0].clear()  # 清空单元格段落数据


if __name__ == '__main__':
    base_path = os.getcwd()
    readfile_path = base_path + '/tihuanmingzi'
    savefile_path = base_path + '/qingxixinxi'
    make_folder(savefile_path)
    for file_name in glob.glob(os.path.join(readfile_path, '*.docx')):
        file = docx.Document(file_name)
        tables = file.tables  # 读取所有表格
        # 遍历所有表，迭代得出每一张表
        for i in range(len(tables)):
            t = tables[i]  # 将每一张表传递给参数t
            p = Process(target=clear_all_user_information(t))  # Process 对象只是一个子任务，运行该任务时系统会自动创建一个子进程
            # 启动一个子进程来运行子任务
            p.start()
            p.join()
            # 子进程完成后，继续运行主进程,保存文件
        if base_path[0] == '/':
            new_name = file_name.strip(readfile_path)  # 去掉文件路径只保留文件名
            file.save(savefile_path + '/' + new_name)  # 保存文件
        else:
            new_name = file_name.split('\\')[-1]  # 去掉文件路径只保留文件名
            file.save(savefile_path + '\\' + new_name)  # 保存文件

