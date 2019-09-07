# coding=utf-8
import docx
from xlsxwriter import *


class DocxToXlsx():
    # 创建类方便调用方法
    def __init__(self, docx_read_path=input('请输入要读取的正确的docx文件路径:'), \
                 docx_write_path=input('请输入要写入的正确的docx文件路径，如果与读取路径重名将覆盖源文件，请慎重:'), \
                 out_xlsx_path=input('请输入正确的输出表格文件路径，保存格式建议为较稳定的xlsx：')):
        self.docx_read_path = docx_read_path
        self.docx_write_path = docx_write_path
        self.out_xlsx_path = out_xlsx_path

    # 完成类的属性定义

    def replace_name(self):
        file = docx.Document(self.docx_read_path)
        tables = file.tables
        name = {}  # 创建字典用于储存替换名字和编号，防止同名同姓
        a = 1  # 创建替换名字的编号
        for table in tables:
            if table.cell(0, 0).text == '编号':  # 查找目标表
                if table.cell(0, 3).text + table.cell(0, 1).text not in name:  # 修改目标表的姓名属性
                    name[table.cell(0, 3).text + table.cell(0, 1).text] = a  # 用字典储存编号
                    a += 1
                    b = str(name.get(table.cell(0, 3).text + table.cell(0, 1).text))  # 将编号转换为字符串
                    table.cell(0, 3).paragraphs[0].clear()
                    table.cell(0, 3).paragraphs[0].add_run(b)  # 用编号替换名字
                    c = '序号'
                    table.cell(0, 2).paragraphs[0].clear()
                    table.cell(0, 2).paragraphs[0].add_run(c)
                else:
                    b = str(name.get(table.cell(0, 3).text + table.cell(0, 1).text))  # 将编号转换为字符串
                    table.cell(0, 3).paragraphs[0].clear()
                    table.cell(0, 3).paragraphs[0].add_run(b)  # 用编号替换名字
                    c = '序号'
                    table.cell(0, 2).paragraphs[0].clear()
                    table.cell(0, 2).paragraphs[0].add_run(c)
        file.save(self.docx_write_path)

    # 替换名字函数完成

    def clear_user_information(self):
        if docx.Document(self.docx_write_path):
            file = docx.Document(self.docx_write_path)
        else:
            file = docx.Document(self.docx_read_path)
        # 防止没进行名字替换而没有写入的docx文件
        tables = file.tables
        for table in tables:
            if table.cell(0, 0).text == '编号':
                m = 0;
                n = 0
                for i in range(len(table.rows)):
                    m += 1  # 统计行数
                    for j in range(len(table.rows[i].cells)):
                        n += 1
            n = int(n / m)  # 统计列数
            for i in range(1, m):
                for j in range(n):
                    for k in range(len(table.cell(i, j).paragraphs)):
                        run = table.cell(i, j).paragraphs[k].clear()  # 清空每个单元格
                    table.cell(1, 0).merge(table.cell(m - 1, n - 1))  # 单元格合并
        file.save(self.docx_write_path)

    # 用户清洗函数完成

    def docx_data_to_xlsx(self):
        workbook = Workbook(self.out_xlsx_path)
        worksheet = workbook.add_worksheet()  # 创建输出表格文件

        def value_clear(dict2={'1': 'a', "2": 's'}):
            key = dict2.keys()
            for i in key:
                dict2[i] = None
            return dict2

        # 定义清空字典值而不改变键的函数

        def strs_are_nums(a='22.22'):
            n = 0
            for i in range(len(a)):
                if '0' <= a[i] <= '9':
                    n += 1
                if a[i] == '.' and '0' <= a[i + 1] <= '9' and '0' <= a[i - 1] <= '9':
                    n += 1
            if n == len(a):
                return True
            else:
                return False

        # 定义判断字符串是否只由浮点数或整数构成的函数

        if docx.Document(self.docx_write_path):
            file = docx.Document(self.docx_write_path)
        else:
            file = docx.Document(self.docx_read_path)
            # 防止没进行名字替换而没有写入的docx文件

        dict0 = {'序号': None, '年龄': None, '性别': None, '编号': None, '内外向(E)': None \
            , '神经质(N)': None, '精神质(P)': None, '掩饰性(L)': None, \
                 '躯体化': None, '强迫症状': None, '人际关系敏感': None, '抑郁': None, '焦虑': None, '敌对': None, '恐怖': None, \
                 '偏执': None, '精神病性': None, '其他': None, '总分': None, '总均分': None, '阳性项目数': None}
        # 创建字典储存标题值

        number = {}
        k = 0;
        a = 0;
        b = 0  # 分别为用户数量，写入文件行数，列数

        for key in dict0.keys():
            worksheet.write(a, b, key)
            b += 1
        a += 1
        # 写入标题行

        tables = file.tables
        for table in tables:
            b = 0
            if table.cell(0, 1).text not in number and strs_are_nums(table.cell(0, 1).text):  # 判断用户
                if k > 1:
                    for value in dict0.values():
                        worksheet.write(a, b, value)
                        b += 1
                    a += 1
                    # 写入上一个用户信息
                    value_clear(dict0)  # 写完后清空字典
                number[table.cell(0, 1).text] = k
                k += 1
                # 利用字典储存用户编号，并且统计用户数量
                for m in range(len(table.rows)):
                    for n in range(len(table.rows[0].cells)):
                        if table.cell(m, n).text in dict0:  # 判断键是否存在在字典中，再写入对应的值
                            if table.cell(m, n + 1).text in ['男', '女'] or strs_are_nums(table.cell(m, n + 1).text):
                                dict0[table.cell(m, n).text] = table.cell(m, n + 1).text
                                # 写入值
                            try:
                                if strs_are_nums(table.cell(m, n + 2).text):
                                    dict0[table.cell(m, n).text] = table.cell(m, n + 2).text
                                # 利用try模块防止列表越界的情况，即n已经代表最后一列时，n+2是不存在的
                            except:
                                pass  # 直接跳过越界
            else:
                for m in range(len(table.rows)):
                    for n in range(len(table.rows[0].cells)):
                        if table.cell(m, n).text in dict0:
                            if table.cell(m, n + 1).text in ['男', '女'] or strs_are_nums(table.cell(m, n + 1).text):
                                dict0[table.cell(m, n).text] = table.cell(m, n + 1).text
                            try:
                                if strs_are_nums(table.cell(m, n + 2).text):
                                    dict0[table.cell(m, n).text] = table.cell(m, n + 2).text
                            except:
                                pass

        for value in dict0.values():
            worksheet.write(a, b, value)
            b += 1
        a += 1
        workbook.close()
# xlsx写入函数完成
