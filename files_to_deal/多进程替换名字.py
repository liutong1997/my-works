# coding = utf-8
from multiprocessing import Process
import docx
import os
import glob
import shutil


# 定义替换名字并记录编号的函数
def replace_name_and_count(t):
    # 引用全局变量
    global a
    global c

    name[t.cell(0, 3).text + t.cell(0, 1).text] = a  # 用字典储存编号
    a += 1  # 编号自增
    b = str(name.get(t.cell(0, 3).text + t.cell(0, 1).text))  # 将编号转换为字符串
    t.cell(0, 3).paragraphs[0].clear()  # 清除名字
    t.cell(0, 3).paragraphs[0].add_run(b)  # 用编号替换名字
    t.cell(0, 2).paragraphs[0].clear()  # 清楚原‘姓名’字符串
    t.cell(0, 2).paragraphs[0].add_run(c)  # 替换为‘编号’


# 定义替换名字的函数
def replace_name(t):
    global c  # 引用全局变量
    c = '序号'  # 赋予c字符串‘序号’

    b = str(name.get(t.cell(0, 3).text + t.cell(0, 1).text))  # 将编号转换为字符串
    t.cell(0, 3).paragraphs[0].clear()  # 清除名字
    t.cell(0, 3).paragraphs[0].add_run(b)  # 用编号替换名字
    t.cell(0, 2).paragraphs[0].clear()  # 清楚原‘姓名’字符串，
    t.cell(0, 2).paragraphs[0].add_run(c)  # 替换为‘编号’


# 定义自动创建文件夹的函数
def make_folder(p):
    if os.path.exists(p):  # 判断文件夹是否存在
        shutil.rmtree(p)  # 如果存在删除原有目录的文件夹以及其中所有文件
    os.mkdir(p)


if __name__ == '__main__':
    # 定义全局变量c,a分别存放‘序号’字符串和
    c = '序号'
    a = 1
    file_path = os.getcwd()  # 获取当前py文件的路径
    # 遍历当前路径的所有.doc文件
    for file_name in glob.glob(os.path.join(file_path, '*.docx')):
        file = docx.Document(file_name)
        tables = file.tables  # 读取所有表格
        name = {}  # 定义字典储存姓名编号
        # 遍历所有表，迭代得出每一张表
        for i in range(len(tables)):
            t = tables[i]  # 将每一张表传递给参数t
            if t.cell(0, 0).text == '编号':  # 查找目标表
                if t.cell(0, 3).text + t.cell(0, 1).text not in name:  # 判断是否已经收录了个人信息
                    p = Process(target=replace_name_and_count(t))  # Process 对象只是一个子任务，运行该任务时系统会自动创建一个子进程
                    # 启动一个子进程来运行子任务
                    p.start()
                    p.join()
                    # 子进程完成后，继续运行主进程,保存文件
                else:
                    p = Process(target=replace_name(t))
                    # 启动一个子进程来运行子任务
                    p.start()
                    p.join()
                    # 子进程完成后，继续运行主进程,保存文件
        new_name = file_name.strip(file_path)
        file.save('new' + new_name)
