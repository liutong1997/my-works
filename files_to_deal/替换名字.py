# coding=utf-8
import docx

file = docx.Document(input('请输入要读取的正确的docx文件路径:'))
tables = file.tables
# 创建字典用于储存替换名字和编号，防止同名同姓
name = {}
# 创建替换名字的编号
a = 1

for table in tables:
    # 查找目标表
    if table.cell(0, 0).text == '编号':
        # 修改目标表的姓名属性
        if table.cell(0, 3).text + table.cell(0, 1).text not in name:
            # 用字典储存编号
            name[table.cell(0, 3).text + table.cell(0, 1).text] = a
            a += 1
            # 将编号转换为字符串
            b = str(name.get(table.cell(0, 3).text + table.cell(0, 1).text))
            table.cell(0, 3).paragraphs[0].clear()
            # 用编号替换名字
            table.cell(0, 3).paragraphs[0].add_run(b)
            c = '序号'
            table.cell(0, 2).paragraphs[0].clear()
            table.cell(0, 2).paragraphs[0].add_run(c)

        else:
            # 将编号转换为字符串
            b = str(name.get(table.cell(0, 3).text + table.cell(0, 1).text))
            table.cell(0, 3).paragraphs[0].clear()
            # 用编号替换名字
            table.cell(0, 3).paragraphs[0].add_run(b)
            c = '序号'
            table.cell(0, 2).paragraphs[0].clear()
            table.cell(0, 2).paragraphs[0].add_run(c)

file.save(input('请输入要写入的正确的docx文件路径，如果与读取路径重名将覆盖源文件，请慎重:'))
